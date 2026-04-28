"""
Montagem do documento Word (.docx) para o relatório de Letras Financeiras.
Formatação: Arial, azul escuro #1F3864, tabelas com linhas alternadas.
"""
import logging
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config.filters import FiltrosConsulta
from config.settings import OUTPUTS_DIR, VERSION
from report.templates import (
    gerar_titulo_relatorio,
    gerar_cabecalho_header,
    gerar_nota_metodologica,
    gerar_sumario_executivo,
    gerar_analise_emissor,
    gerar_contexto_macro,
    gerar_consideracoes_finais,
)
from report.tables import (
    preparar_tabela_emissoes,
    preparar_tabela_resumo_emissor,
    preparar_tabela_por_tipo,
)
from utils.helpers import formatar_data, timestamp_arquivo

logger = logging.getLogger(__name__)

# Paleta de cores
COR_AZUL_ESCURO = RGBColor(0x1F, 0x38, 0x64)
COR_AZUL_CLARO = RGBColor(0xDD, 0xE8, 0xF5)
COR_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
COR_CINZA_CLARO = RGBColor(0xF2, 0xF2, 0xF2)
COR_TEXTO = RGBColor(0x26, 0x26, 0x26)


def _set_cell_bg(cell, cor_hex: str) -> None:
    """Define cor de fundo de uma célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    tcPr.append(shd)


def _set_cell_borders(cell) -> None:
    """Define bordas sutis em uma célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _paragrafo_titulo(doc: Document, texto: str, nivel: int = 1) -> None:
    """Adiciona parágrafo de título estilizado."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    run.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = COR_AZUL_ESCURO
    if nivel == 1:
        run.font.size = Pt(16)
    elif nivel == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    p.space_before = Pt(12 if nivel <= 2 else 6)
    p.space_after = Pt(4)


def _paragrafo_corpo(doc: Document, texto: str, negrito: bool = False) -> None:
    """Adiciona parágrafo de corpo de texto."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(texto)
    run.bold = negrito
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = COR_TEXTO
    p.space_after = Pt(6)


def _adicionar_tabela(
    doc: Document,
    cabecalhos: list[str],
    linhas: list[dict],
    larguras: list[float] | None = None,
) -> None:
    """
    Adiciona tabela formatada ao documento.
    Linha de cabeçalho: azul escuro com texto branco.
    Linhas de dados: alternadas (branco / cinza claro).
    """
    if not linhas:
        _paragrafo_corpo(doc, "Nenhum dado disponível para esta tabela.")
        return

    n_cols = len(cabecalhos)
    tabela = doc.add_table(rows=1, cols=n_cols)
    tabela.alignment = WD_TABLE_ALIGNMENT.LEFT
    tabela.style = "Table Grid"

    # Cabeçalho
    row_header = tabela.rows[0]
    for i, cab in enumerate(cabecalhos):
        cell = row_header.cells[i]
        cell.text = cab
        _set_cell_bg(cell, "1F3864")
        _set_cell_borders(cell)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run(cab)
        run.text = cab
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = COR_BRANCO

    # Larguras de coluna
    if larguras and len(larguras) == n_cols:
        for i, larg in enumerate(larguras):
            for row in tabela.rows:
                row.cells[i].width = Inches(larg)

    # Linhas de dados
    for idx, linha in enumerate(linhas):
        row = tabela.add_row()
        cor_bg = "FFFFFF" if idx % 2 == 0 else "F2F2F2"
        for i, cab in enumerate(cabecalhos):
            cell = row.cells[i]
            valor = str(linha.get(cab, ""))
            cell.text = valor
            _set_cell_bg(cell, cor_bg)
            _set_cell_borders(cell)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if para.runs:
                run = para.runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.color.rgb = COR_TEXTO

    doc.add_paragraph()  # Espaço após tabela


def _adicionar_header_footer(doc: Document, filtros: FiltrosConsulta) -> None:
    """Adiciona cabeçalho e rodapé ao documento."""
    secao = doc.sections[0]

    # Cabeçalho
    header = secao.header
    header.is_linked_to_previous = False
    p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p_header.clear()
    run = p_header.add_run(gerar_cabecalho_header(filtros))
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = COR_AZUL_ESCURO
    run.bold = True
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Rodapé
    footer = secao.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.clear()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run(
        f"Agente B3 v{VERSION} | Gerado em {date.today().strftime('%d/%m/%Y')} | "
        "Uso restrito — não constitui recomendação de investimento"
    )
    run_f.font.name = "Arial"
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def gerar_docx(
    df: pd.DataFrame,
    filtros: FiltrosConsulta,
    fontes_usadas: list[str],
    total_antes_filtro: int,
    output_path: Path | None = None,
) -> Path:
    """
    Monta e salva o relatório Word (.docx).

    Retorna o caminho do arquivo gerado.
    """
    doc = Document()

    # Configurar margens
    secao = doc.sections[0]
    secao.left_margin = Cm(2.5)
    secao.right_margin = Cm(2.5)
    secao.top_margin = Cm(2.5)
    secao.bottom_margin = Cm(2.0)

    # Cabeçalho/rodapé
    _adicionar_header_footer(doc, filtros)

    # ── CAPA ─────────────────────────────────────────────────────────────────
    p_capa = doc.add_paragraph()
    p_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_capa.space_before = Pt(40)
    titulo_completo = gerar_titulo_relatorio(filtros)
    for linha in titulo_completo.split("\n"):
        run = p_capa.add_run(linha + "\n")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(18 if "LETRAS" in linha else 13)
        run.font.color.rgb = COR_AZUL_ESCURO

    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_data = p_data.add_run(f"Data de geração: {date.today().strftime('%d/%m/%Y')}")
    run_data.font.name = "Arial"
    run_data.font.size = Pt(10)
    run_data.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_page_break()

    # ── SUMÁRIO EXECUTIVO ────────────────────────────────────────────────────
    texto_sumario = gerar_sumario_executivo(df, filtros, total_antes_filtro)
    for linha in texto_sumario.split("\n\n"):
        if linha.startswith("SUMÁRIO"):
            _paragrafo_titulo(doc, linha, nivel=1)
        else:
            _paragrafo_corpo(doc, linha)

    # ── CONTEXTO MACROECONÔMICO ──────────────────────────────────────────────
    doc.add_paragraph()
    texto_macro = gerar_contexto_macro()
    for linha in texto_macro.split("\n\n"):
        if linha.startswith("CONTEXTO"):
            _paragrafo_titulo(doc, linha, nivel=1)
        else:
            _paragrafo_corpo(doc, linha)

    # ── TABELA RESUMO POR EMISSOR ────────────────────────────────────────────
    doc.add_paragraph()
    _paragrafo_titulo(doc, "RESUMO POR EMISSOR", nivel=1)
    linhas_resumo = preparar_tabela_resumo_emissor(df)
    _adicionar_tabela(
        doc,
        ["Emissor", "Rating", "Nº Emissões", "Total Captado", "Spread Médio"],
        linhas_resumo,
        larguras=[2.5, 0.8, 0.9, 1.3, 1.8],
    )

    # ── TABELA COMPARATIVA POR TIPO DE LF ────────────────────────────────────
    _paragrafo_titulo(doc, "SPREADS POR TIPO DE LF", nivel=1)
    linhas_tipo = preparar_tabela_por_tipo(df)
    _adicionar_tabela(
        doc,
        ["Tipo de LF", "Nº Emissões", "Spread Mín.", "Spread Máx.", "Spread Médio", "Volume Total"],
        linhas_tipo,
        larguras=[2.0, 0.9, 1.2, 1.2, 1.2, 1.3],
    )

    # ── ANÁLISE INDIVIDUAL POR EMISSOR ───────────────────────────────────────
    doc.add_paragraph()
    _paragrafo_titulo(doc, "ANÁLISE POR EMISSOR", nivel=1)

    emissores = df["emissor"].unique() if not df.empty else []
    for emissor in sorted(emissores):
        subset = df[df["emissor"] == emissor]
        rating = subset["rating_nota_normalizado"].iloc[0] or subset["rating_nota"].iloc[0] if len(subset) > 0 else "N/D"
        emissoes_list = []
        for _, row in subset.iterrows():
            emissoes_list.append({
                "tipo_lf": row.get("tipo_lf", "senior"),
                "valor_emissao": row.get("valor_emissao"),
                "data_emissao": row.get("data_emissao"),
                "indexador": row.get("indexador", "CDI"),
                "spread_percentual": row.get("spread_percentual"),
                "prazo_label": row.get("prazo_label", "N/D"),
            })

        texto_emissor = gerar_analise_emissor({
            "nome": emissor,
            "rating": rating,
            "emissoes": emissoes_list,
        })

        for bloco in texto_emissor.split("\n\n"):
            if "—" in bloco and "Rating" in bloco:
                _paragrafo_titulo(doc, bloco, nivel=2)
            elif bloco.strip():
                _paragrafo_corpo(doc, bloco)

    # ── TABELA DETALHADA DE EMISSÕES ─────────────────────────────────────────
    doc.add_page_break()
    _paragrafo_titulo(doc, "DETALHAMENTO DAS EMISSÕES", nivel=1)
    linhas_emissoes = preparar_tabela_emissoes(df)
    _adicionar_tabela(
        doc,
        ["Emissor", "Rating", "Tipo", "Data Emissão", "Vencimento", "Valor (R$ mi)", "Taxa", "Prazo", "Status"],
        linhas_emissoes,
        larguras=[1.8, 0.7, 1.5, 1.0, 1.0, 1.0, 1.5, 0.8, 0.8],
    )

    # ── CONSIDERAÇÕES FINAIS ─────────────────────────────────────────────────
    doc.add_paragraph()
    texto_final = gerar_consideracoes_finais(df, filtros)
    for linha in texto_final.split("\n\n"):
        if linha.startswith("CONSIDERAÇÕES"):
            _paragrafo_titulo(doc, linha, nivel=1)
        else:
            _paragrafo_corpo(doc, linha)

    # ── NOTA METODOLÓGICA ────────────────────────────────────────────────────
    doc.add_paragraph()
    texto_metodologia = gerar_nota_metodologica(
        filtros, fontes_usadas, len(df), total_antes_filtro
    )
    for linha in texto_metodologia.split("\n\n"):
        if linha.startswith("NOTA"):
            _paragrafo_titulo(doc, linha, nivel=1)
        else:
            _paragrafo_corpo(doc, linha)

    # ── SALVAR ───────────────────────────────────────────────────────────────
    if output_path is None:
        output_dir = Path(filtros.output_dir) if filtros.output_dir else OUTPUTS_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        nome_arquivo = (
            f"relatorio_lf_rating_{filtros.descricao_rating().replace(' a ', '_')}_"
            f"{timestamp_arquivo()}.docx"
        )
        output_path = output_dir / nome_arquivo

    doc.save(str(output_path))
    logger.info(f"Relatório gerado: {output_path}")
    return output_path
